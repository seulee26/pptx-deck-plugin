"""Parse an Excel/Word/PDF input into a synthesized outline.json for slide-matcher.

v0.3 — ADAPTIVE SUMMARIZATION
    An input file is summarized into an executive-ready deck (target 12–25
    slides). Each xlsx sheet becomes ONE section with ≤ 3 items. For each
    sheet, the parser:

      1. Strips the banner title / blank separator rows to locate the real
         header row.
      2. Detects the sheet's *shape*:
           - WIDE: every row is a complete entity (e.g. 거래처 list, 프로젝트
             portfolio). Picks the most informative cols + top N rows ranked
             by the biggest numeric col.
           - LONG: first col is an entity name that repeats across rows with a
             metric name in col 2 (e.g. 사업부별 P&L). Pivots LONG→WIDE by
             `entity × selected-metrics`.
           - BLOCKED: sheet is split into sub-tables by `[A] / [B]` labels.
             Each block is re-run through the WIDE/LONG detector.
      3. Produces up to `MAX_ITEMS_PER_SECTION` items, each trimmed to fit
         the biggest native template table (9 rows × 5 cols).

Size caps
    MAX_SECTIONS          = 7    overflow sheets merged into "기타"
    MAX_ITEMS_PER_SECTION = 3
    MAX_TABLE_ROWS        = 8    (fits slide 56's 9×5 after +1 header)
    MAX_TABLE_COLS        = 5
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

MARKER_TYPES = {"kpi", "process", "matrix", "table", "image", "content"}

MAX_SECTIONS = 7
MAX_ITEMS_PER_SECTION = 3
MAX_TABLE_ROWS = 5   # fits slide 54/55 (6 rows incl. header) cleanly
MAX_TABLE_COLS = 5

KPI_HINT_WORDS = ("pnl", "p&l", "손익", "매출", "kpi", "스코어", "파이프라인", "실적", "bsc")
KPI_LABEL_HINTS = ("매출액", "영업이익", "영업이익률", "수주", "달성률", "점수", "합계", "연간합계", "가중")

# Columns to surface when pivoting "사업부 × 항목" long → wide P&L.
PNL_PRIMARY_METRICS = ("매출액", "영업이익", "영업이익률", "매출총이익(GP)", "GP마진율")

# Columns worth keeping for wide transactional tables, in priority order.
SUMMARY_COL_HINTS = (
    "합계", "연간합계", "연간", "총계", "총액", "누적",
    "매출액", "매출", "영업이익", "영업이익률",
    "달성률", "진행률", "수주확률", "수금률",
    "상태", "등급", "결과", "산업군", "사업부", "담당",
    "연봉", "인건비", "인원",
    "가중점수", "가중금액",
    "Next Action", "핵심이슈",
)

YEAR_PREFIX_RE = re.compile(r"^\d{4}[_\- ]*")


# ───────────────────────────── utilities ────────────────────────────────────

def clean_heading(name: str) -> str:
    n = YEAR_PREFIX_RE.sub("", name).replace("_", " ").strip()
    return n or name


def fmt_cell(v) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        if v.is_integer():
            return f"{int(v):,}" if abs(v) >= 1000 else str(int(v))
        # Small decimals → percent-ish display when between -1..1.
        if -1 < v < 1 and v != 0:
            return f"{v * 100:.1f}%"
        return f"{v:,.1f}".rstrip("0").rstrip(".")
    s = str(v).strip()
    # Apply thousands separator to plain integers.
    if s.replace("-", "").isdigit() and len(s) >= 4:
        try:
            return f"{int(s):,}"
        except ValueError:
            pass
    return s


def is_short_label(s: str) -> bool:
    return bool(s) and len(s) <= 24 and not any(ch.isdigit() for ch in s[:1])


def looks_numeric(s: str) -> bool:
    if not s:
        return False
    x = s.replace(",", "").replace("%", "").replace(" ", "")
    try:
        float(x)
        return True
    except ValueError:
        return False


def find_header_row(rows: list[list[str]]) -> int:
    for i, row in enumerate(rows[:8]):
        short = sum(1 for c in row if is_short_label(c))
        if short >= 3 and not any(looks_numeric(c) for c in row[:3]):
            return i
    return 0


def strip_trailing_blanks(header: list[str], body: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    while header and not header[-1]:
        header.pop()
        for row in body:
            if len(row) > len(header):
                row.pop()
    return header, body


# ───────────────────────────── shape detection ──────────────────────────────

def detect_long_pivot_cols(rows: list[list[str]]) -> bool:
    """Return True if the sheet looks like an entity × metric LONG table
    (first col repeats with blanks; second col is a short-label metric)."""
    if not rows or len(rows) < 4:
        return False
    first_col_values = [r[0] for r in rows if r]
    nonempty = [v for v in first_col_values if v]
    # Repetition signal: many blanks in col 0 relative to non-blanks.
    return (len(first_col_values) - len(nonempty)) >= len(nonempty)


def pivot_long_to_wide(
    header: list[str], body: list[list[str]]
) -> tuple[list[str], list[list[str]]]:
    """Pivot an entity × metric LONG table to WIDE using PNL_PRIMARY_METRICS."""
    if len(header) < 3:
        return header, body

    entity_col = 0
    metric_col = 1
    # Numeric/summary columns are everything after col 2 that carries a total.
    # Prefer 연간합계 / 전년실적 / YoY증감률 if present.
    summary_col_idx = None
    for i, h in enumerate(header):
        if any(tok in h for tok in ("연간합계", "합계", "전년실적", "연간")):
            summary_col_idx = i
            break
    if summary_col_idx is None:
        # fallback: last col
        summary_col_idx = len(header) - 1

    # Collect entity → {metric: value} map.
    current_entity = ""
    entity_rows: dict[str, dict[str, str]] = {}
    entity_order: list[str] = []
    for row in body:
        if not row:
            continue
        entity = row[entity_col].strip() if entity_col < len(row) else ""
        metric = row[metric_col].strip() if metric_col < len(row) else ""
        value = row[summary_col_idx].strip() if summary_col_idx < len(row) else ""
        if entity:
            current_entity = entity
            if current_entity not in entity_rows:
                entity_rows[current_entity] = {}
                entity_order.append(current_entity)
        if current_entity and metric and metric not in entity_rows[current_entity]:
            entity_rows[current_entity][metric] = value

    if not entity_order:
        return header, body

    # Pick the primary metrics (filter to ones that actually appear).
    observed = {m for e in entity_order for m in entity_rows[e]}
    chosen_metrics = [m for m in PNL_PRIMARY_METRICS if m in observed]
    if len(chosen_metrics) < 2:
        # fall back to any metrics (preserve order of first-seen)
        seen_order: list[str] = []
        for e in entity_order:
            for m in entity_rows[e]:
                if m not in seen_order:
                    seen_order.append(m)
        chosen_metrics = seen_order[:4]

    chosen_metrics = chosen_metrics[: MAX_TABLE_COLS - 1]

    new_header = [header[entity_col]] + chosen_metrics
    new_body = []
    for e in entity_order:
        row = [e] + [entity_rows[e].get(m, "") for m in chosen_metrics]
        new_body.append(row)

    return new_header, new_body


def select_wide_columns(
    header: list[str], body: list[list[str]]
) -> tuple[list[str], list[list[str]]]:
    """Keep ≤ MAX_TABLE_COLS cols: the ID col + the most informative summary
    cols by SUMMARY_COL_HINTS priority."""
    if len(header) <= MAX_TABLE_COLS:
        return header, body

    keep_idx = [0]  # ID / name column
    # Rank remaining cols by hint priority.
    scored: list[tuple[int, int]] = []
    for i, h in enumerate(header[1:], start=1):
        score = -1
        for prio, hint in enumerate(SUMMARY_COL_HINTS):
            if hint in h:
                score = 1000 - prio
                break
        scored.append((score, i))
    scored.sort(key=lambda x: (-x[0], x[1]))
    for _, idx in scored:
        if len(keep_idx) >= MAX_TABLE_COLS:
            break
        if idx not in keep_idx:
            keep_idx.append(idx)
    keep_idx.sort()

    new_header = [header[i] for i in keep_idx]
    new_body = [[row[i] if i < len(row) else "" for i in keep_idx] for row in body]
    return new_header, new_body


def rank_rows_by_magnitude(
    header: list[str], body: list[list[str]], max_rows: int = MAX_TABLE_ROWS
) -> list[list[str]]:
    """Pick top `max_rows` rows by the biggest numeric column (ignoring the
    first/ID column). Falls back to keeping order if no numeric col."""
    if len(body) <= max_rows:
        return body

    # Choose the column with the most numeric cells.
    col_numeric_count = [0] * len(header)
    for row in body:
        for i, cell in enumerate(row):
            if i < len(col_numeric_count) and looks_numeric(cell):
                col_numeric_count[i] += 1
    # Skip col 0 (ID) when ranking.
    best_col = None
    best_count = 0
    for i in range(1, len(col_numeric_count)):
        if col_numeric_count[i] > best_count:
            best_count = col_numeric_count[i]
            best_col = i
    if best_col is None:
        return body[:max_rows]

    def to_num(s: str) -> float:
        try:
            return float(s.replace(",", "").replace("%", ""))
        except ValueError:
            return float("-inf")

    ranked = sorted(body, key=lambda r: to_num(r[best_col]) if best_col < len(r) else float("-inf"), reverse=True)
    return ranked[:max_rows]


TOTAL_ROW_TOKENS = ("합계", "소계", "총계", "총합", "total", "subtotal", "grand")


def prune_blank_body_rows(body: list[list[str]]) -> list[list[str]]:
    pruned = []
    for row in body:
        non_empty = sum(1 for c in row if c.strip())
        if non_empty < 2:
            continue
        # Drop rows whose label is a total/subtotal — they distort ranking
        # and duplicate the section's takeaway.
        label = (row[0] or "").strip().lower() if row else ""
        if any(tok in label for tok in TOTAL_ROW_TOKENS):
            continue
        pruned.append(row)
    return pruned


def build_table_item(
    title: str, header: list[str], body: list[list[str]]
) -> dict:
    header, body = strip_trailing_blanks(header, body)
    body = prune_blank_body_rows(body)

    if detect_long_pivot_cols(body):
        header, body = pivot_long_to_wide(header, body)

    header, body = select_wide_columns(header, body)
    body = rank_rows_by_magnitude(header, body, MAX_TABLE_ROWS)

    return {
        "type": "table",
        "title": title,
        "body": "",
        "data": {"headers": header, "rows": body},
    }


# ───────────────────────────── block splitter ───────────────────────────────

def split_by_bracket_labels(
    rows: list[list[str]],
) -> list[tuple[str, list[list[str]]]]:
    blocks: list[tuple[str, list[list[str]]]] = []
    current_label = ""
    current_rows: list[list[str]] = []
    bracket_re = re.compile(r"^\[[A-Za-z0-9]\]\s*(.+)$")
    for row in rows:
        first = row[0] if row else ""
        m = bracket_re.match(first)
        if m and not any(row[1:]):
            if current_rows:
                blocks.append((current_label, current_rows))
            current_label = m.group(1).strip()
            current_rows = []
        else:
            current_rows.append(row)
    if current_rows:
        blocks.append((current_label, current_rows))
    return blocks


# ───────────────────────────── xlsx parsing ─────────────────────────────────

def parse_marker_row(rows: list[list[str]], i: int) -> tuple[dict, int]:
    row = rows[i]
    kind = row[0].lower()
    title = row[1] if len(row) > 1 else ""
    body = row[2] if len(row) > 2 else ""
    data: dict = {}

    if kind == "kpi":
        data = {"value": row[3] if len(row) > 3 else "", "delta": row[4] if len(row) > 4 else ""}
        return {"type": kind, "title": title, "body": body, "data": data}, i + 1
    if kind == "process":
        data = {"steps": [c for c in row[3:] if c]}
        return {"type": kind, "title": title, "body": body, "data": data}, i + 1
    if kind == "matrix":
        data = {
            "S": row[3] if len(row) > 3 else "",
            "W": row[4] if len(row) > 4 else "",
            "O": row[5] if len(row) > 5 else "",
            "T": row[6] if len(row) > 6 else "",
        }
        return {"type": kind, "title": title, "body": body, "data": data}, i + 1
    if kind == "table":
        headers = [h for h in row[3:] if h]
        tbl_rows: list[list[str]] = []
        j = i + 1
        while j < len(rows) and rows[j] and rows[j][0].lower() not in MARKER_TYPES:
            tbl_rows.append([c for c in rows[j] if c])
            j += 1
        return build_table_item(title or "Table", headers, tbl_rows), j
    return {"type": kind, "title": title, "body": body, "data": {}}, i + 1


def synthesize_sheet(sheet_name: str, rows: list[list[str]]) -> dict:
    heading = clean_heading(sheet_name)

    has_markers = any(row and row[0].lower() in MARKER_TYPES for row in rows[:30])
    if has_markers:
        items: list[dict] = []
        i = 0
        while i < len(rows):
            row = rows[i]
            marker = row[0].lower() if row else ""
            if marker in MARKER_TYPES:
                item, i = parse_marker_row(rows, i)
                items.append(item)
            else:
                i += 1
        return {"heading": heading, "items": items[:MAX_ITEMS_PER_SECTION]}

    blocks = split_by_bracket_labels(rows)
    items: list[dict] = []

    blocks.sort(key=lambda b: len(b[1]), reverse=True)
    for block_label, block_rows in blocks[:MAX_ITEMS_PER_SECTION]:
        hdr_idx = find_header_row(block_rows)
        header = [h for h in block_rows[hdr_idx] if h]
        body_rows = [
            [fmt_cell(c) for c in r]
            for r in block_rows[hdr_idx + 1 :]
            if any(c for c in r)
        ]
        if not header or not body_rows:
            continue
        title = block_label or heading
        items.append(build_table_item(title, header, body_rows))

    if not items:
        joined = "\n".join(
            " ".join(fmt_cell(c) for c in r if c) for r in rows[:8] if any(c for c in r)
        )
        items = [{"type": "content", "title": heading, "body": joined[:300], "data": {}}]

    return {"heading": heading, "items": items[:MAX_ITEMS_PER_SECTION]}


def parse_xlsx(path: Path) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    meta: dict[str, str] = {}
    sections: list[dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [[fmt_cell(c) for c in row] for row in ws.iter_rows(values_only=True)]
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            continue

        if sheet_name.lower() == "meta":
            for row in rows:
                if len(row) >= 2 and row[0]:
                    meta[row[0].lower()] = row[1]
            continue

        sections.append(synthesize_sheet(sheet_name, rows))

    if len(sections) > MAX_SECTIONS:
        overflow = sections[MAX_SECTIONS - 1 :]
        merged_items: list[dict] = []
        for s in overflow:
            merged_items.extend(s["items"])
        sections = sections[: MAX_SECTIONS - 1]
        sections.append({"heading": "기타", "items": merged_items[:MAX_ITEMS_PER_SECTION]})

    title = meta.get("title") or clean_heading(path.stem)
    return {
        "title": title,
        "subtitle": meta.get("subtitle", ""),
        "company": meta.get("company", ""),
        "sections": sections,
        "closing": meta.get("closing", "Thank you"),
    }


# ───────────────────────────── docx / pdf ───────────────────────────────────

def parse_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(path)
    title = ""
    sections: list[dict] = []
    current_section: dict | None = None
    current_item: dict | None = None

    def flush_item() -> None:
        nonlocal current_item
        if current_item and current_section is not None:
            current_section["items"].append(current_item)
        current_item = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading 1" in style and not title:
            title = text
        elif "heading 2" in style:
            flush_item()
            current_section = {"heading": text, "items": []}
            sections.append(current_section)
        elif "heading 3" in style:
            flush_item()
            current_item = {"type": "content", "title": text, "body": "", "data": {}}
        else:
            if current_section is None:
                current_section = {"heading": "Overview", "items": []}
                sections.append(current_section)
            if current_item is None:
                current_item = {"type": "content", "title": "", "body": "", "data": {}}
            current_item["body"] = (current_item["body"] + "\n" + text).strip()

            if any(ch.isdigit() for ch in text) and len(text) < 40 and any(s in text for s in ["%", "K", "M", "B"]):
                current_item["type"] = "kpi"
                current_item["data"] = {"value": text, "delta": ""}
    flush_item()

    for tbl in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if not rows:
            continue
        item = build_table_item(rows[0][0] if rows[0] else "Table", rows[0], rows[1:])
        if not sections:
            sections.append({"heading": "Data", "items": []})
        sections[-1]["items"].append(item)

    for s in sections:
        s["items"] = s["items"][:MAX_ITEMS_PER_SECTION]
    sections = sections[:MAX_SECTIONS]

    return {
        "title": title or path.stem,
        "subtitle": "",
        "company": "",
        "sections": sections,
        "closing": "Thank you",
    }


def parse_pdf(path: Path) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    title = path.stem
    if pages:
        first_line = next((ln.strip() for ln in pages[0].splitlines() if ln.strip()), "")
        if first_line:
            title = first_line[:80]

    sections: list[dict] = []
    for text in pages[1:]:
        if len(sections) >= MAX_SECTIONS:
            break
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if not lines:
            continue
        heading = lines[0][:80]
        body = "\n".join(lines[1:])[:400]
        sections.append(
            {
                "heading": heading,
                "items": [{"type": "content", "title": heading, "body": body, "data": {}}],
            }
        )

    if not sections:
        sections = [
            {
                "heading": "Content",
                "items": [{"type": "content", "title": title, "body": "\n".join(pages)[:400], "data": {}}],
            }
        ]

    return {"title": title, "subtitle": "", "company": "", "sections": sections, "closing": "Thank you"}


def parse(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return parse_xlsx(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    raise ValueError(f"unsupported input type: {suffix}")


def main() -> None:
    if len(sys.argv) != 3:
        sys.exit("usage: parse_input.py <input> <outline.json>")
    src = Path(sys.argv[1]).expanduser().resolve()
    dst = Path(sys.argv[2]).expanduser().resolve()
    dst.parent.mkdir(parents=True, exist_ok=True)

    try:
        outline = parse(src)
    except Exception as e:
        outline = {
            "title": src.stem,
            "subtitle": "",
            "company": "",
            "sections": [
                {
                    "heading": "Content",
                    "items": [{"type": "content", "title": src.stem, "body": f"(parse failed: {e})", "data": {}}],
                }
            ],
            "closing": "Thank you",
        }

    dst.write_text(json.dumps(outline, ensure_ascii=False, indent=2))
    n_sections = len(outline["sections"])
    n_items = sum(len(s["items"]) for s in outline["sections"])
    print(f"outline written: {dst}")
    print(f"  sections: {n_sections}  items: {n_items}")


if __name__ == "__main__":
    main()
